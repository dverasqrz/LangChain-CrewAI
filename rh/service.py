from __future__ import annotations

import json
import logging
import os
import re
import time
import uuid
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from pypdf import PdfReader
from langchain_core.prompts import ChatPromptTemplate

from rh.config import (
    ConfigError,
    AppConfig,
    LLMConfig,
    get_app_config,
    get_llm_config,
    validate_provider_environment,
)

logger = logging.getLogger(__name__)

SUPPORTED_CV_EXTENSIONS = {".pdf", ".txt", ".docx"}
DEFAULT_ANALYSIS_FILE = "rh_curriculos.json"


class RHServiceError(Exception):
    """Erro de negócio ou execução do serviço de RH."""


@dataclass(frozen=True)
class RHAnalysisResult:
    request_id: str
    candidate_name: str
    job_title: str
    provider: str
    model: str
    alignment_score: float
    summarization: str
    details: dict[str, Any]
    prompt_preview: str
    elapsed_ms: int


class RHService:
    def __init__(self, app_config: AppConfig | None = None) -> None:
        self.app_config = app_config or get_app_config()
        self._llm = None

    @property
    def llm(self):
        if self._llm is None:
            self._llm = self._build_llm(self.app_config.ai_provider)
        return self._llm

    def _build_llm(self, provider: str):
        config = get_llm_config(provider)
        validate_provider_environment(provider)

        try:
            if config.provider == "groq":
                from langchain_groq import ChatGroq

                return ChatGroq(
                    model=config.model,
                    temperature=config.temperature,
                    max_tokens=None,
                    timeout=config.timeout_seconds,
                    max_retries=config.max_retries,
                )

            if config.provider == "openai":
                from langchain_openai import ChatOpenAI

                return ChatOpenAI(
                    model=config.model,
                    temperature=config.temperature,
                    timeout=config.timeout_seconds,
                    max_retries=config.max_retries,
                )

            raise RHServiceError(f"Provider não suportado: {provider}.")
        except ConfigError as exc:
            raise RHServiceError(f"Erro de configuração: {exc}") from exc
        except Exception as exc:
            raise RHServiceError(
                f"Falha ao inicializar o modelo do provider '{provider}': {exc}"
            ) from exc

    @staticmethod
    def _extract_text_from_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text_parts: list[str] = []
        for page in reader.pages:
            text_parts.append(page.extract_text() or "")
        return "\n\n".join(text_parts).strip()

    @staticmethod
    def _extract_text_from_txt(file_path: str) -> str:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as handle:
            return handle.read().strip()

    @staticmethod
    def _extract_text_from_docx(file_path: str) -> str:
        try:
            from docx import Document
        except ImportError:
            raise RHServiceError(
                "O módulo 'python-docx' não está instalado. "
                "Execute: pip install python-docx"
            )
        
        document = Document(file_path)
        text_parts: list[str] = []
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        for table in document.tables:
            for row in table.rows:
                row_text = " ".join(cell.text for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)
        return "\n\n".join(text_parts).strip()

    def parse_resume(self, file_path: str) -> str:
        extension = Path(file_path).suffix.lower()
        if extension not in SUPPORTED_CV_EXTENSIONS:
            raise RHServiceError(
                f"Formato de currículo não suportado: {extension}. Use PDF, DOCX ou TXT."
            )

        if extension == ".pdf":
            return self._extract_text_from_pdf(file_path)
        if extension == ".docx":
            return self._extract_text_from_docx(file_path)
        return self._extract_text_from_txt(file_path)

    @staticmethod
    def _build_prompt_schema() -> str:
        return json.dumps(
            {
                "candidate_name": "Nome completo do candidato",
                "summary": "Resumo de perfil e experiência relevante",
                "experience_overview": "Resumo de experiência profissional, áreas de atuação e resultados",
                "skills": "Lista de habilidades e competências técnicas e comportamentais",
                "education": "Formação acadêmica e cursos relevantes",
                "strengths": "Pontos fortes alinhados com a vaga",
                "development_opportunities": "Principais lacunas e pontos de atenção",
                "alignment_score": "Nota de alinhamento entre 0.0 e 10.0",
                "recommendations": "Recomendações práticas para o próximo passo",
            },
            ensure_ascii=False,
            indent=2,
        )

    def _build_prompt(self, job_text: str, resume_text: str) -> ChatPromptTemplate:
        system_prompt = self.app_config.default_system_prompt
        template = (
            "Sistema: {system_prompt}\n\n"
            "Analise o currículo do candidato e compare com a vaga abaixo. "
            "Retorne apenas um JSON válido contendo os campos definidos no schema. "
            "Não adicione texto extra ou explicações fora do JSON.\n\n"
            "Vaga:\n{job_text}\n\n"
            "Currículo:\n{resume_text}\n\n"
            "Schema desejado:\n{schema}\n\n"
            "Para o campo alignment_score, responda apenas com um número entre 0.0 e 10.0, "
            "usando uma única casa decimal quando possível."
        )

        return ChatPromptTemplate.from_template(template)

    @staticmethod
    def _extract_json(text: str) -> dict[str, Any] | None:
        if "{" not in text or "}" not in text:
            return None

        start = text.find("{")
        end = text.rfind("}") + 1
        candidate = text[start:end].strip()
        try:
            return json.loads(candidate)
        except json.JSONDecodeError:
            # remove markdown fences and stray annotations
            candidate = re.sub(r"```(?:json)?", "", candidate)
            start = candidate.find("{")
            end = candidate.rfind("}") + 1
            if start == -1 or end == 0:
                return None
            try:
                return json.loads(candidate[start:end])
            except json.JSONDecodeError:
                return None

    def analyze_candidate(
        self,
        job_title: str,
        job_description: str,
        job_details: str,
        resume_text: str,
    ) -> RHAnalysisResult:
        if not resume_text:
            raise RHServiceError("O currículo não contém texto processável.")

        job_text = f"Título: {job_title}\nDescrição: {job_description}\nDetalhes: {job_details}".strip()
        template_text = (
            "Sistema: {system_prompt}\n\n"
            "Analise o currículo do candidato e compare com a vaga abaixo. "
            "Retorne apenas um JSON válido contendo os campos definidos no schema. "
            "Não adicione texto extra ou explicações fora do JSON.\n\n"
            "Vaga:\n{job_text}\n\n"
            "Currículo:\n{resume_text}\n\n"
            "Schema desejado:\n{schema}\n\n"
            "Para o campo alignment_score, responda apenas com um número entre 0.0 e 10.0, "
            "usando uma única casa decimal quando possível."
        )
        prompt_template = ChatPromptTemplate.from_template(template_text)
        start = uuid.uuid4().hex
        chain = prompt_template | self.llm
        started = time.perf_counter()

        output = chain.invoke(
            {
                "system_prompt": self.app_config.default_system_prompt,
                "job_text": job_text,
                "resume_text": resume_text,
                "schema": self._build_prompt_schema(),
            }
        )

        elapsed_ms = int((time.perf_counter() - started) * 1000)
        model_response = output.content.strip()
        parsed = self._extract_json(model_response)

        if not parsed:
            raise RHServiceError("Não foi possível interpretar a resposta do modelo como JSON.")

        candidate_name = parsed.get("candidate_name") or parsed.get("name") or "Candidato não identificado"
        alignment_score = float(parsed.get("alignment_score") or parsed.get("score") or 0.0)

        if alignment_score < 0.0:
            alignment_score = 0.0
        if alignment_score > 10.0:
            alignment_score = 10.0

        return RHAnalysisResult(
            request_id=start,
            candidate_name=str(candidate_name).strip(),
            job_title=job_title,
            provider=self.app_config.ai_provider,
            model=self.llm.model,
            alignment_score=alignment_score,
            summarization=str(parsed.get("summary") or parsed.get("experience_overview") or "").strip(),
            details={
                "experience_overview": parsed.get("experience_overview", ""),
                "skills": parsed.get("skills", []),
                "education": parsed.get("education", ""),
                "strengths": parsed.get("strengths", []),
                "development_opportunities": parsed.get("development_opportunities", []),
                "recommendations": parsed.get("recommendations", ""),
            },
            prompt_preview=template_text,
            elapsed_ms=elapsed_ms,
        )

    @staticmethod
    def load_analyses(path_json: str | Path = DEFAULT_ANALYSIS_FILE) -> list[dict[str, Any]]:
        path = Path(path_json)
        if not path.exists():
            return []
        with path.open("r", encoding="utf-8") as handle:
            return json.load(handle)

    @staticmethod
    def save_analysis(
        analysis: RHAnalysisResult,
        path_json: str | Path = DEFAULT_ANALYSIS_FILE,
    ) -> None:
        path = Path(path_json)
        path.parent.mkdir(parents=True, exist_ok=True)
        existing = []
        if path.exists():
            with path.open("r", encoding="utf-8") as handle:
                existing = json.load(handle)

        candidate_names = {entry.get("candidate_name") for entry in existing if isinstance(entry, dict)}
        if analysis.candidate_name in candidate_names:
            existing = [entry for entry in existing if entry.get("candidate_name") != analysis.candidate_name]

        existing.append(
            {
                "request_id": analysis.request_id,
                "candidate_name": analysis.candidate_name,
                "job_title": analysis.job_title,
                "provider": analysis.provider,
                "model": analysis.model,
                "alignment_score": analysis.alignment_score,
                "summarization": analysis.summarization,
                "details": analysis.details,
            }
        )

        with path.open("w", encoding="utf-8") as handle:
            json.dump(existing, handle, indent=2, ensure_ascii=False)

    @staticmethod
    def get_top_matches(
        analyses: list[dict[str, Any]],
        top_n: int = 5,
    ) -> list[dict[str, Any]]:
        return sorted(
            analyses,
            key=lambda item: float(item.get("alignment_score", 0.0) or 0.0),
            reverse=True,
        )[:top_n]
